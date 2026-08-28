from __future__ import annotations

import os
from html import escape
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "deliverables"
DELIVERABLES.mkdir(parents=True, exist_ok=True)

PDF_OUTPUT_PATH = DELIVERABLES / "Bloodchain_Technical_Architecture_Overview.pdf"
DOCX_OUTPUT_PATH = DELIVERABLES / "Bloodchain_Technical_Architecture_Overview.docx"
DOC_DATE = "August 2026"
DOC_REF = "BC-ARCH-2026-V2.4"

# Palette
COLOR_INK = colors.HexColor("#0B1220")
COLOR_SLATE = colors.HexColor("#334155")
COLOR_MUTED = colors.HexColor("#64748B")
COLOR_CRIMSON = colors.HexColor("#991B1B")
COLOR_CRIMSON_DARK = colors.HexColor("#7F1D1D")
COLOR_CYAN = colors.HexColor("#0284C7")
COLOR_CYAN_BG = colors.HexColor("#F0F9FF")
COLOR_CRIMSON_BG = colors.HexColor("#FEF2F2")
COLOR_CARD_BG = colors.HexColor("#F8FAFC")
COLOR_LINE = colors.HexColor("#E2E8F0")
COLOR_ACCENT = colors.HexColor("#0D9488")


def get_styles():
    base = getSampleStyleSheet()
    return {
        "doc_title": ParagraphStyle(
            "BC_DocTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=COLOR_INK,
            alignment=TA_LEFT,
            spaceAfter=10,
        ),
        "doc_subtitle": ParagraphStyle(
            "BC_DocSubtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=COLOR_SLATE,
            spaceAfter=14,
        ),
        "h1": ParagraphStyle(
            "BC_H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=18,
            textColor=COLOR_CRIMSON_DARK,
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "BC_H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=15,
            textColor=COLOR_INK,
            spaceBefore=9,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "BC_H3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=9.5,
            leading=13,
            textColor=COLOR_CYAN,
            spaceBefore=5,
            spaceAfter=3,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "BC_Body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=13.5,
            textColor=COLOR_SLATE,
            spaceAfter=6,
        ),
        "body_bold": ParagraphStyle(
            "BC_BodyBold",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            leading=13.5,
            textColor=COLOR_INK,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "BC_Bullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=12.5,
            textColor=COLOR_SLATE,
            leftIndent=12,
            firstLineIndent=-8,
            spaceAfter=3,
        ),
        "table_header": ParagraphStyle(
            "BC_TableHeader",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=10.5,
            textColor=colors.white,
        ),
        "table_cell": ParagraphStyle(
            "BC_TableCell",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=7.8,
            leading=10.5,
            textColor=COLOR_INK,
        ),
        "callout_text": ParagraphStyle(
            "BC_CalloutText",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=12.5,
            textColor=COLOR_SLATE,
        ),
        "tag": ParagraphStyle(
            "BC_Tag",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=8,
            leading=10,
            textColor=COLOR_CRIMSON,
            alignment=TA_LEFT,
        ),
    }


STYLES = get_styles()


def p(text: str, style_name: str = "body"):
    return Paragraph(text, STYLES[style_name])


def info_box(title: str, body: str, bg_color=COLOR_CARD_BG, border_color=COLOR_LINE):
    content = [
        Paragraph(f"<b>{escape(title).upper()}</b>", STYLES["tag"]),
        Spacer(1, 0.12 * cm),
        Paragraph(body, STYLES["callout_text"]),
    ]
    t = Table([[content]], colWidths=[17.1 * cm])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), bg_color),
                ("BOX", (0, 0), (-1, -1), 0.75, border_color),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return t


def data_table(headers: list[str], rows: list[list[str]], widths: list[float]):
    data = [[Paragraph(f"<b>{escape(h)}</b>", STYLES["table_header"]) for h in headers]]
    for row in rows:
        data.append(
            [
                Paragraph(escape(str(val)).replace("\n", "<br/>"), STYLES["table_cell"])
                for val in row
            ]
        )
    t = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), COLOR_INK),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.3, COLOR_LINE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, COLOR_CARD_BG]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return t


def page_header_footer(canvas, doc):
    canvas.saveState()
    w, h = A4
    canvas.setFillColor(COLOR_CRIMSON_DARK)
    canvas.rect(0, h - 0.4 * cm, w, 0.4 * cm, fill=1, stroke=0)
    
    canvas.setFillColor(COLOR_MUTED)
    canvas.setFont("Helvetica-Bold", 7)
    canvas.drawString(1.5 * cm, h - 0.75 * cm, "BLOODCHAIN  |  TECHNICAL ARCHITECTURE & PLATFORM DOSSIER")
    canvas.drawRightString(w - 1.5 * cm, h - 0.75 * cm, f"DOC REF: {DOC_REF}")
    
    canvas.setStrokeColor(COLOR_LINE)
    canvas.setLineWidth(0.4)
    canvas.line(1.5 * cm, h - 0.9 * cm, w - 1.5 * cm, h - 0.9 * cm)

    canvas.line(1.5 * cm, 0.9 * cm, w - 1.5 * cm, 0.9 * cm)
    canvas.setFont("Helvetica", 7)
    canvas.drawString(1.5 * cm, 0.55 * cm, f"https://bloodchain.life  •  UniPod Botswana  •  {DOC_DATE}")
    canvas.drawRightString(w - 1.5 * cm, 0.55 * cm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    doc = SimpleDocTemplate(
        str(PDF_OUTPUT_PATH),
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.3 * cm,
        bottomMargin=1.3 * cm,
    )

    story = []

    # COVER BLOCK
    story.append(Spacer(1, 0.3 * cm))
    story.append(p("SOVEREIGN HEALTH INFRASTRUCTURE", "tag"))
    story.append(p("Bloodchain: National Blood Supply Grid & Cryptographic Custody Architecture", "doc_title"))
    story.append(
        p(
            "An enterprise technical overview of Botswana's decentralised blood supply coordination "
            "platform—connecting voluntary donors, laboratory screening, hospital cold storage, and national crisis routing.",
            "doc_subtitle",
        )
    )

    meta_headers = ["Document Reference", "Platform URL", "Engineering & Origin", "Target Sector"]
    meta_rows = [
        [
            f"{DOC_REF}\n({DOC_DATE})",
            "https://bloodchain.life\nAPI: bloodchain-api.onrender.com",
            "Gift Jr Nakedi\nUniPod, University of Botswana",
            "National Blood Transfusion Service (NBTS)\nMinistry of Health, Hospital Networks",
        ]
    ]
    story.append(data_table(meta_headers, meta_rows, [3.5 * cm, 4.8 * cm, 4.4 * cm, 4.4 * cm]))
    story.append(Spacer(1, 0.4 * cm))

    story.append(
        info_box(
            "EXECUTIVE PLATFORM SUMMARY",
            "<b>Bloodchain</b> solves critical blood stock shortages, cold-chain expiration wastage, and delayed maternal postpartum "
            "hemorrhage response through a unified 4-pillar software suite. Engineered around the <b>Strand Ledger Protocol</b> "
            "(tamper-proof cryptographic custody) and the <b>Torrent Routing Engine</b> (geo-spatial algorithmic blood balancing), "
            "the platform delivers real-time vein-to-vein traceability from voluntary donation to bedside transfusion.",
            bg_color=COLOR_CYAN_BG,
            border_color=COLOR_CYAN,
        )
    )
    story.append(Spacer(1, 0.4 * cm))

    # SECTION 1
    story.append(p("1. National Problem Framing & Strategic Objective", "h1"))
    story.append(
        p(
            "Sub-Saharan national blood services face severe structural supply bottlenecks. In Botswana, the National Blood Transfusion Service (NBTS) "
            "targets an annual collection threshold of <b>35,000+ units</b>, yet regional facilities face acute deficits. Maternal hemorrhage and acute trauma "
            "require immediate blood matching within 30 minutes, yet manual telephone coordination and unlinked hospital reserves lead to critical delays."
        )
    )

    problem_headers = ["Structural Challenge", "Current Legacy Reality", "Bloodchain Solution Architecture"]
    problem_rows = [
        [
            "Maternal Deficits & Acute Trauma",
            "Severe postpartum hemorrhage requires blood within 30 minutes; manual phone coordination causes lethal delays.",
            "Instant multi-facility shortage alerts dispatch targeted notifications to matching universal donors (O- / O+) via Pulse Mobile.",
        ],
        [
            "Cold-Chain Expiration & Wastage",
            "Strict 35-42 day red blood cell shelf-life; peripheral clinics lose valuable units due to lack of grid visibility.",
            "Torrent routing algorithm automatically flags nearing-expiry units and coordinates automated rebalancing to high-volume trauma hubs.",
        ],
        [
            "Verification & Audit Integrity",
            "Paper logs create chain-of-custody blind spots during inter-hospital transit and serological clearance.",
            "Strand cryptographic ledger seals each step (collection, viral screening, quarantine clearance, cross-match, transfusion) in SHA-256 hashes.",
        ],
    ]
    story.append(data_table(problem_headers, problem_rows, [3.8 * cm, 6.2 * cm, 7.1 * cm]))
    story.append(Spacer(1, 0.4 * cm))

    # SECTION 2
    story.append(p("2. Four-Pillar Unified Product Ecosystem", "h1"))
    story.append(
        p(
            "Bloodchain operates as a role-isolated monorepo architecture, providing dedicated interfaces tailored to each stakeholder:"
        )
    )

    app_headers = ["Application", "Target Audience", "Deployment Mode", "Core Capabilities"]
    app_rows = [
        [
            "Pulse\n(Web & Mobile)",
            "Voluntary Donors & General Public",
            "Android APK (Expo EAS)\nWeb App (React 19 / Vite)",
            "• Digital Sovereign Donor ID with cryptographic QR card\n• Live national reserve deficit telemetry & shortage alerts\n• Airtime reward tracking & donation history verification\n• Seamless lookup / sign-in & consent management",
        ],
        [
            "Sanctum",
            "Hospital Blood Bank Staff & Clinicians",
            "Web Terminal\n(React 19 / Vite)",
            "• Real-time cold storage unit custody tracking\n• Cross-matching and emergency reservation\n• Inter-facility transfer initiation and confirmation\n• Bedside unit release and transfusion logging",
        ],
        [
            "Crucible",
            "Serology & Pathology Lab Technicians",
            "Lab Workstation\n(React 19 / Vite)",
            "• Mandatory viral screening workflows (HIV, HBV, HCV, Syphilis)\n• ABO & Rh blood typing validation\n• Quarantine release authority and test certification\n• Batch cryptographic signature generation",
        ],
        [
            "Vigil",
            "NBTS Command, MoH & National Leadership",
            "Command Center\n(React 19 / Vite)",
            "• National live inventory grid by hospital & blood type\n• Real-time critical shortage heat maps and predictive deficits\n• National transit ledger and inter-district transfer logs\n• Executive crisis response override tools",
        ],
    ]
    story.append(data_table(app_headers, app_rows, [2.4 * cm, 3.2 * cm, 3.8 * cm, 7.7 * cm]))
    story.append(Spacer(1, 0.4 * cm))

    # SECTION 3
    story.append(p("3. Technical Architecture & Monorepo Topology", "h1"))
    story.append(
        p(
            "Bloodchain is engineered with a contract-first monorepo pattern using <b>pnpm workspaces</b>, "
            "ensuring strict type synchronization between the central database, microservices, and client frontends:"
        )
    )

    tech_headers = ["Architecture Layer", "Technologies", "Role & Engineering Implementation"]
    tech_rows = [
        [
            "Frontend Applications",
            "React 19, React Native, Expo Router v54, Vite 8, Tailwind CSS",
            "Role-specific responsive web dashboards and compiled native Android APK with offline token caching and biometric/QR card export.",
        ],
        [
            "API Microservices",
            "Node.js 24, Express 5, TypeScript 5.9, Zod v4",
            "High-throughput RESTful microservice handling identity, unit lifecycle state transitions, transfer dispatch, and telemetry.",
        ],
        [
            "Shared API Spec",
            "OpenAPI 3.0, Orval Codegen, React Query v5",
            "Single source of truth contract (`lib/api-spec/openapi.yaml`). Generates typed client SDKs and server validation schemas automatically.",
        ],
        [
            "Database & Ledger",
            "PostgreSQL 16, Drizzle ORM, Node Crypto (SHA-256)",
            "Relational data store enforcing foreign key constraints, automated migrations, seed orchestrations, and immutable event hashing.",
        ],
        [
            "Cloud Infrastructure",
            "Render Cloud Platform, Expo EAS, Custom SSL Domain",
            "Production API (`bloodchain-api.onrender.com`), edge CDN web deployments on `bloodchain.life`, and cloud binary builder.",
        ],
    ]
    story.append(data_table(tech_headers, tech_rows, [3.2 * cm, 4.3 * cm, 9.6 * cm]))
    story.append(Spacer(1, 0.4 * cm))

    # SECTION 4
    story.append(p("4. Proprietary Core Engines: Strand & Torrent", "h1"))
    story.append(
        info_box(
            "ENGINE 1: THE STRAND CRYPTOGRAPHIC LEDGER",
            "<b>The Strand Protocol</b> guarantees that no blood unit can be transfused without full compliance verification. "
            "Every state transition generates an immutable SHA-256 hash incorporating the previous block hash, timestamp, technician identity, "
            "temperature log, and lab screening results: <br/>"
            "<code>Block Hash = SHA256(PrevHash + UnitID + EventType + LabResults + FacilityID + Timestamp)</code><br/>"
            "Spanning: <b>Donation Enrolment → Collection → Viral Clearance → Cold Storage → Transfer Transit → Transfusion</b>.",
            bg_color=COLOR_CARD_BG,
            border_color=COLOR_LINE,
        )
    )
    story.append(Spacer(1, 0.25 * cm))

    story.append(
        info_box(
            "ENGINE 2: THE TORRENT SMART ROUTING ALGORITHM",
            "<b>The Torrent Engine</b> optimizes national blood allocation by evaluating four dynamic parameters in real time:<br/>"
            "1. <b>Shelf-Life Expiry Risk:</b> Prioritizes units nearing the 35-day window for immediate dispatch to high-consumption trauma centers.<br/>"
            "2. <b>Blood Group Compatibility Matrix:</b> Protects rare universal reserves (O-negative) by strictly enforcing Rh-iso matching.<br/>"
            "3. <b>Transit Matrix & Cold-Chain Constraints:</b> Computes transit time corridors (e.g., PMH Gaborone to Nyangabgwe or Maun General).<br/>"
            "4. <b>Deficit Severity Score:</b> Weights maternity wards and trauma surgeries at 3x standard replenishment priority.",
            bg_color=COLOR_CARD_BG,
            border_color=COLOR_LINE,
        )
    )
    story.append(Spacer(1, 0.4 * cm))

    # SECTION 5
    story.append(p("5. Data Governance, Consent, & Security Compliance", "h1"))
    story.append(
        p(
            "• <b>Donor Pseudonymization:</b> National Omang ID numbers and medical histories are decoupled from public QR codes. Public scans display only cryptographic validation tokens and ABO/Rh groups.",
            "bullet",
        )
    )
    story.append(
        p(
            "• <b>Explicit Applied Telemetry Consent:</b> The Pulse registration engine incorporates explicit legal consent clauses covering algorithmic testing, shortage notification dispatch, and NBTS clinical reporting.",
            "bullet",
        )
    )
    story.append(
        p(
            "• <b>Role-Based Access Control (RBAC):</b> Laboratory release overrides require certified Crucible technician credentials, preventing unverified or seropositive blood units from entering active hospital inventories.",
            "bullet",
        )
    )
    story.append(
        p(
            "• <b>Sovereign Hosting & Isolation:</b> All health records and ledger events remain under sovereign jurisdiction, with zero dependencies on foreign ad networks or untrusted analytics trackers.",
            "bullet",
        )
    )
    story.append(Spacer(1, 0.4 * cm))

    # SECTION 6
    story.append(p("6. Implementation Roadmap & Institutional Deployment", "h1"))
    roadmap_headers = ["Phase", "Milestone & Deployment Scope", "Target Infrastructure", "Status"]
    roadmap_rows = [
        [
            "Phase 1",
            "Platform Engineering, API Monorepo, & Simulation Grid",
            "Live on https://bloodchain.life (Render + EAS)",
            "COMPLETED\n(Production Live)",
        ],
        [
            "Phase 2",
            "NBTS Gaborone & Princess Marina Hospital (PMH) Pilot",
            "Crucible lab terminal + Sanctum blood bank integration",
            "READY FOR DEPLOYMENT",
        ],
        [
            "Phase 3",
            "National Transit Corridor Integration (Nyangabgwe & Maun)",
            "Torrent inter-facility dispatch & cold-chain logger sync",
            "SCHEDULED (Q4 2026)",
        ],
        [
            "Phase 4",
            "USSD / SMS Offline Fallback & Africa's Talking Gateway",
            "National network coverage for non-smartphone donors",
            "PLANNED (Q1 2027)",
        ],
    ]
    story.append(data_table(roadmap_headers, roadmap_rows, [2.0 * cm, 8.0 * cm, 4.5 * cm, 2.6 * cm]))
    story.append(Spacer(1, 0.5 * cm))

    story.append(
        info_box(
            "INSTITUTIONAL CONTACT & INQUIRIES",
            "<b>Bloodchain Sovereign Health Technologies</b><br/>"
            "Lead Engineer / Founder: <b>Gift Jr Nakedi</b><br/>"
            "Incubation & Engineering Facility: <b>UniPod, University of Botswana, Gaborone</b><br/>"
            "Platform URL: <b>https://bloodchain.life</b>  •  Contact: <b>giftjrnakedi@gmail.com</b>",
            bg_color=COLOR_CRIMSON_BG,
            border_color=COLOR_CRIMSON,
        )
    )

    doc.build(story, onFirstPage=page_header_footer, onLaterPages=page_header_footer)
    print(f"Created PDF: {PDF_OUTPUT_PATH}")


def build_docx():
    doc = Document()

    p_title = doc.add_heading("BLOODCHAIN", level=0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_sub = doc.add_paragraph("National Blood Supply Grid & Cryptographic Custody Architecture")
    p_sub.runs[0].font.size = Pt(13)
    p_sub.runs[0].font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
    p_sub.runs[0].font.bold = True

    p_meta = doc.add_paragraph(
        f"Document Reference: {DOC_REF}  |  Date: {DOC_DATE}  |  Platform: https://bloodchain.life\n"
        "Engineered by: Gift Jr Nakedi (UniPod, University of Botswana)"
    )
    p_meta.runs[0].font.size = Pt(9.5)
    p_meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Bloodchain is Botswana's sovereign digital health coordination infrastructure designed to eliminate critical "
        "blood shortages, prevent cold-chain expiration, and accelerate maternal postpartum hemorrhage emergency response. "
        "It connects voluntary blood donors, hospital blood banks, screening laboratories, and the Ministry of Health "
        "through a unified, real-time cryptographic supply grid."
    )

    doc.add_heading("2. Four-Pillar Product Ecosystem", level=1)
    p_apps = [
        ("Pulse (Mobile & Web)", "Voluntary donor engagement, digital sovereign donor card, live deficit warnings, and airtime reward tracking."),
        ("Sanctum (Hospital Terminal)", "Hospital blood bank inventory management, cold-chain monitoring, cross-matching, and transfusion logging."),
        ("Crucible (Lab Workstation)", "Mandatory viral screening (HIV, HBV, HCV, Syphilis), ABO/Rh validation, quarantine clearance, and batch cryptographic certification."),
        ("Vigil (National Command Grid)", "National real-time inventory overview for NBTS and Ministry of Health, shortage prediction heat maps, and transit rebalancing controls."),
    ]
    for app_name, desc in p_apps:
        p = doc.add_paragraph()
        r1 = p.add_run(f"• {app_name}: ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.add_run(desc)

    doc.add_heading("3. Core Engines: Strand & Torrent", level=1)
    doc.add_paragraph(
        "• Strand Cryptographic Ledger: Enforces immutable SHA-256 hash chains across every state transition "
        "(Donation → Collection → Viral Screening → Storage → Transfer → Transfusion), preventing tampering and counterfeit logging."
    )
    doc.add_paragraph(
        "• Torrent Smart Routing Engine: Evaluates shelf-life expiration risks, blood group compatibility (O- universal allocation), "
        "transit times across Botswana corridors, and emergency maternity ward deficits to automatically rebalance national stocks."
    )

    doc.add_heading("4. Technology Stack & Deployment", level=1)
    doc.add_paragraph(
        "• Monorepo Architecture: pnpm workspaces with React 19, React Native (Expo Router v54), Express 5 API Server, "
        "PostgreSQL 16, Drizzle ORM, and OpenAPI v3 typed contracts.\n"
        "• Production Endpoints: Live at https://bloodchain.life with API services hosted on Render (bloodchain-api.onrender.com) "
        "and compiled native Android binaries via Expo EAS."
    )

    doc.add_heading("5. Implementation Roadmap", level=1)
    doc.add_paragraph(
        "• Phase 1: National Simulation Grid & Monorepo Platform (Completed & Live on bloodchain.life)\n"
        "• Phase 2: Pilot Deployment at Princess Marina Hospital & NBTS Gaborone (Ready for Launch)\n"
        "• Phase 3: Regional expansion to Nyangabgwe (Francistown) & Maun General Hospital (Q4 2026)\n"
        "• Phase 4: USSD / SMS national fallback integration via Africa's Talking (Q1 2027)"
    )

    doc.save(str(DOCX_OUTPUT_PATH))
    print(f"Created DOCX: {DOCX_OUTPUT_PATH}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
