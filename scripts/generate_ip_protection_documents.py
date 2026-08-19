from __future__ import annotations

import os
from html import escape
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "deliverables"
REPORT_PATH = DELIVERABLES / "Bloodchain_IP_Protection_Codebase_Report.pdf"
LETTER_PATH = DELIVERABLES / "Bloodchain_NBTS_Introduction_Letter_Draft.pdf"
LETTER_DOCX_PATH = DELIVERABLES / "Bloodchain_NBTS_Introduction_Letter_Draft.docx"
PREPARED_DATE = "19 August 2026"

INK = colors.HexColor("#0B1220")
SLATE = colors.HexColor("#44546A")
MUTED = colors.HexColor("#6B7280")
CRIMSON = colors.HexColor("#B4233C")
CRIMSON_DARK = colors.HexColor("#7A1025")
CYAN = colors.HexColor("#007A99")
PALE_CRIMSON = colors.HexColor("#FCE8EC")
PALE_CYAN = colors.HexColor("#E7F7FB")
PALE_GRAY = colors.HexColor("#F4F6F8")
LINE = colors.HexColor("#D8DEE6")


def build_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "BC_Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            leading=33,
            textColor=INK,
            alignment=TA_LEFT,
            spaceAfter=14,
        ),
        "subtitle": ParagraphStyle(
            "BC_Subtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=12,
            leading=17,
            textColor=SLATE,
            spaceAfter=16,
        ),
        "h1": ParagraphStyle(
            "BC_H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            textColor=CRIMSON_DARK,
            spaceBefore=14,
            spaceAfter=9,
        ),
        "h2": ParagraphStyle(
            "BC_H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=17,
            textColor=INK,
            spaceBefore=10,
            spaceAfter=5,
        ),
        "h3": ParagraphStyle(
            "BC_H3",
            parent=base["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=14,
            textColor=CRIMSON_DARK,
            spaceBefore=7,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "BC_Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=13.2,
            textColor=INK,
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "BC_Small",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.6,
            leading=10,
            textColor=SLATE,
            spaceAfter=4,
        ),
        "tiny": ParagraphStyle(
            "BC_Tiny",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=6.7,
            leading=8.2,
            textColor=SLATE,
        ),
        "callout": ParagraphStyle(
            "BC_Callout",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=9.2,
            leading=13.2,
            textColor=CRIMSON_DARK,
            leftIndent=8,
            rightIndent=8,
            spaceBefore=3,
            spaceAfter=3,
        ),
        "code": ParagraphStyle(
            "BC_Code",
            parent=base["Code"],
            fontName="Courier",
            fontSize=7.1,
            leading=9.3,
            textColor=INK,
            leftIndent=6,
            spaceBefore=2,
            spaceAfter=2,
        ),
        "letter": ParagraphStyle(
            "BC_Letter",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.4,
            leading=15.2,
            textColor=INK,
            spaceAfter=10,
        ),
        "letter_title": ParagraphStyle(
            "BC_LetterTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=25,
            textColor=INK,
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
    }


STYLES = build_styles()


def p(text: str, style: str = "body") -> Paragraph:
    return Paragraph(text, STYLES[style])


def bullet_list(items, level=0):
    # Do not use ListFlowable here: in the current ReportLab runtime its
    # bullet token is extracted/rendered as the literal word "bullet".
    # A two-column table gives us a stable, real bullet glyph and keeps
    # long evidence paths aligned without changing the body text.
    body_style = STYLES["body"]
    bullet_style = ParagraphStyle(
        f"BC_BulletGlyph_{level}",
        parent=body_style,
        fontName="Helvetica",
        fontSize=10,
        leading=13.2,
        alignment=TA_CENTER,
        textColor=CRIMSON,
    )
    rows = [
        [Paragraph("•", bullet_style), Paragraph(item, body_style)]
        for item in items
    ]
    table = Table(
        rows,
        colWidths=[0.45 * cm, 16.65 * cm],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, -1), 0),
                ("RIGHTPADDING", (0, 0), (0, -1), 5),
                ("LEFTPADDING", (1, 0), (1, -1), 0),
                ("RIGHTPADDING", (1, 0), (1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def info_box(title: str, content: str, accent=PALE_CRIMSON):
    table = Table(
        [[p(title, "h3")], [p(content, "body")]],
        colWidths=[17.1 * cm],
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), accent),
                ("BOX", (0, 0), (-1, -1), 0.6, LINE),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, 0), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    return table


def data_table(headers, rows, widths, font_size=7.1):
    data = [
        [Paragraph(f"<b>{escape(str(h))}</b>", STYLES["small"]) for h in headers]
    ]
    for row in rows:
        data.append(
            [
                Paragraph(escape(str(value)).replace("\n", "<br/>"), STYLES["tiny"])
                for value in row
            ]
        )
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), INK),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.3, LINE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, PALE_GRAY]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def report_header(canvas, doc):
    canvas.saveState()
    width, height = A4
    canvas.setFillColor(INK)
    canvas.rect(0, height - 0.7 * cm, width, 0.7 * cm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 7.5)
    canvas.drawString(1.55 * cm, height - 0.44 * cm, "BLOODCHAIN  |  CONFIDENTIAL INTERNAL IP INVENTORY")
    canvas.setFillColor(MUTED)
    canvas.setFont("Helvetica", 7)
    canvas.drawRightString(width - 1.55 * cm, 0.58 * cm, f"Prepared {PREPARED_DATE}  •  Page {doc.page}")
    canvas.setStrokeColor(LINE)
    canvas.line(1.55 * cm, 0.82 * cm, width - 1.55 * cm, 0.82 * cm)
    canvas.restoreState()


def letter_header(canvas, doc):
    canvas.saveState()
    width, height = A4
    canvas.setFillColor(CRIMSON_DARK)
    canvas.rect(0, height - 0.6 * cm, width, 0.6 * cm, fill=1, stroke=0)
    canvas.setFillColor(MUTED)
    canvas.setFont("Helvetica", 7)
    canvas.drawRightString(width - 1.55 * cm, 0.58 * cm, f"Draft for review • {PREPARED_DATE} • Page {doc.page}")
    canvas.restoreState()


def cover(story):
    story.extend(
        [
            Spacer(1, 3.0 * cm),
            p("BLOODCHAIN", "h3"),
            p("Confidential Codebase & Intellectual Property Protection Report", "title"),
            p(
                "A detailed technical inventory of the Bloodchain platform, its product applications, "
                "shared services, authored workflows, evidence locations, and third-party boundaries.",
                "subtitle",
            ),
            Spacer(1, 0.4 * cm),
            info_box(
                "DOCUMENT STATUS",
                "<b>Confidential — internal use only.</b> Prepared for founder, investor and institutional "
                "pitch preparation. This report records the current workspace snapshot as of "
                f"<b>{PREPARED_DATE}</b>; it is not a legal opinion, patentability assessment, "
                "copyright registration, or substitute for advice from Botswana-qualified intellectual-property "
                "and data-protection counsel.",
            ),
            Spacer(1, 0.45 * cm),
            data_table(
                ["Prepared for", "Purpose", "Evidence basis"],
                [
                    (
                        "Bloodchain leadership",
                        "Preserve and communicate what has been created before external pitching",
                        "Current source tree, package manifests, API contract, database schema, artifact configuration and generated deliverables",
                    )
                ],
                [4.0 * cm, 6.0 * cm, 7.1 * cm],
            ),
            Spacer(1, 1.2 * cm),
            p(
                "Working position: Bloodchain is a Botswana-focused, multi-role blood-supply coordination "
                "platform. Its strongest authored assets are its domain model, operational workflows, "
                "laboratory state controls, chain-of-custody implementation, API contract, original product "
                "composition and Bloodchain brand expression—not the underlying open-source frameworks.",
                "callout",
            ),
            Spacer(1, 4.5 * cm),
            p("Prepared from the Bloodchain repository workspace", "small"),
            p("Document reference: BC-IP-TECH-2026-08-19 • Version 1.0", "small"),
            PageBreak(),
        ]
    )


def add_report_intro(story):
    story.extend(
        [
            p("1. Purpose, scope and handling", "h1"),
            p(
                "This document is a technical and product-origin inventory for internal IP protection. It identifies "
                "the software, workflows, content, visual expression and evidence locations that together make up "
                "Bloodchain. It is designed to help the team describe its work consistently during investor, "
                "government and health-sector discussions; preserve a record of pre-disclosure development; and "
                "separate proprietary contribution from third-party software.",
                "body",
            ),
            info_box(
                "HOW TO USE THIS REPORT",
                "Keep this PDF with its source repository export, Git history, dated build artifacts, contributor "
                "records, design files and signed confidentiality/assignment records. Before disclosing detailed "
                "architecture, share only the level of detail appropriate to the audience and use an NDA or "
                "institutional confidentiality arrangement where suitable.",
                PALE_CYAN,
            ),
            p("Included", "h2"),
            bullet_list(
                [
                    "All registered Bloodchain artifacts: public demo, web donor portal, mobile donor app, hospital terminal, national dashboard, laboratory workstation, API service, pitch deck and internal mockup sandbox.",
                    "Shared API specification, React Query client, server-side Zod validators, database schema, seed/demo data and repository tooling.",
                    "Product flows, endpoints, data entities, laboratory rules, evidence paths and an authored-versus-generated/third-party boundary.",
                ]
            ),
            p("Not included or not asserted", "h2"),
            bullet_list(
                [
                    "Any personal secrets, credentials, private keys, database connection values or user data.",
                    "Ownership of React, Vite, Expo, PostgreSQL, Drizzle, Zod, Orval, Tailwind, Radix/shadcn-derived primitives, Replit services, EAS or other third-party/open-source dependencies.",
                    "A claim that the current demonstration is clinically certified, production authorised, integrated with laboratory instruments/EHRs, or compliant with every applicable health-data requirement.",
                ]
            ),
            p("2. Executive asset register", "h1"),
            data_table(
                ["Asset / code package", "Audience and purpose", "Primary delivery surface", "Authored value represented"],
                [
                    ("Bloodchain Demo\n@workspace/demo", "Investors, partners and public-facing narrative", "Web root /", "Brand story, national-system explanation, product launcher, presentation content"),
                    ("Pulse Web\n@workspace/pulse", "Blood donors", "Web /pulse/", "Registration, donor dashboard, history, donor card, facilities and guidance"),
                    ("Pulse Mobile\n@workspace/pulse-mobile", "Blood donors on mobile", "Expo / Android build", "Native donor registration/card workflow, persistent donor context, shareable card export"),
                    ("Sanctum\n@workspace/sanctum", "Hospital blood-bank staff", "Web /sanctum/", "Facility inventory, unit logging, transfer routing and ledger views"),
                    ("Vigil\n@workspace/vigil", "National/ministry operations", "Web /vigil/", "National telemetry, facility visibility, shortage alerts and routing tool"),
                    ("Crucible\n@workspace/crucible", "Blood-processing laboratory teams", "Web /crucible/", "Screening, component separation, state transitions and chain-of-custody ledger"),
                    ("API Server\n@workspace/api-server", "All Bloodchain applications", "API /api", "Shared business API, routing logic, laboratory guards and service orchestration"),
                    ("Bloodchain Pitch\n@workspace/bloodchain-pitch", "Pitch meetings", "Web /bloodchain-pitch/", "Ten-slide narrative, presentation viewer and slide interaction model"),
                    ("Mockup Sandbox\n@workspace/mockup-sandbox", "Internal design workflow", "Internal web artifact", "Component-preview infrastructure; supporting tooling rather than a deployed product surface"),
                ],
                [3.3 * cm, 4.4 * cm, 3.0 * cm, 6.4 * cm],
            ),
            PageBreak(),
        ]
    )


def add_architecture(story):
    story.extend(
        [
            p("3. Platform architecture and original system design", "h1"),
            p(
                "Bloodchain is implemented as a pnpm/TypeScript monorepo. Multiple role-specific user interfaces "
                "consume a common API contract and a common PostgreSQL-backed operational model. The architecture "
                "keeps domain rules in the API/database layer while giving each user group a distinct, tailored "
                "surface. This is a material product-design decision: donor engagement, hospital operations, national "
                "coordination and laboratory processing are represented as connected, but separate, workflows.",
                "body",
            ),
            data_table(
                ["Layer", "Source of truth", "What it contains", "IP / evidence note"],
                [
                    ("Product applications", "artifacts/*", "React/Vite web applications, Expo mobile application and authored slide deck", "Custom pages, flows, narrative, branding and composition are the relevant authored expression"),
                    ("API contract", "lib/api-spec/openapi.yaml", "Public operations, request/response schemas and laboratory data types", "Primary evidence of the platform interface and intended system behavior"),
                    ("Server implementation", "artifacts/api-server/src", "Express routes, facility/routing behavior, lab service and transactional mutations", "Primary evidence for custom business logic"),
                    ("Data model", "lib/db/src/schema", "Donor, blood-unit, facility, laboratory processing, component and event tables", "Primary evidence for the domain model and lifecycle"),
                    ("Shared client/validation", "lib/api-client-react and lib/api-zod", "Generated React Query hooks and generated Zod validators from OpenAPI", "Generated outputs are derivative artifacts; retain source OpenAPI as the authored master"),
                    ("Deployment/tooling", "artifact TOMLs, root manifests, scripts", "Artifact routes, builds, workspace configuration and post-merge script", "Operational evidence; do not confuse platform configuration with sole product IP"),
                ],
                [2.7 * cm, 4.0 * cm, 5.6 * cm, 4.8 * cm],
            ),
            p("Role-aware product system", "h2"),
            bullet_list(
                [
                    "<b>Donor side:</b> Pulse web and Pulse Mobile present registration, eligibility and personal donation history in donor-friendly language and visuals.",
                    "<b>Facility side:</b> Sanctum gives hospital staff inventory, unit intake, inter-facility transfer and ledger views.",
                    "<b>National side:</b> Vigil aggregates stock, shortage and facility information into a ministry/command-centre interface with refresh cadence and routing support.",
                    "<b>Laboratory side:</b> Crucible establishes a dedicated processing workflow from test intake through screening, component separation, quarantine/release/discard, with a chronological hashed lab event trail.",
                    "<b>Shared operational core:</b> the API and database bring those surfaces together around donors, facilities, blood units and laboratory processing records.",
                ]
            ),
            p("Authoring and generation boundary", "h2"),
            p(
                "The authored OpenAPI document drives code generation into the React client and Zod validator packages "
                "through Orval. Treat <font name='Courier'>lib/api-spec/openapi.yaml</font> and the authored server/database "
                "implementations as the stronger evidence of original work. Treat files under "
                "<font name='Courier'>src/generated</font>, compiled <font name='Courier'>dist</font> output and "
                "vendor packages as generated or third-party material unless separately modified and documented.",
                "body",
            ),
            PageBreak(),
        ]
    )


def app_profile(story, number, name, route, audience, purpose, contains, flows, data, evidence, ip_note):
    story.extend(
        [
            p(f"{number}. {name}", "h1"),
            data_table(
                ["Route / runtime", "Audience", "Product role"],
                [(route, audience, purpose)],
                [4.2 * cm, 4.6 * cm, 8.3 * cm],
            ),
            p("What this application contains", "h2"),
            bullet_list(contains),
            p("Principal user flows", "h2"),
            bullet_list(flows),
            p("Data/API relationship", "h2"),
            p(data, "body"),
            p("Evidence locations", "h2"),
            bullet_list([f"<font name='Courier'>{escape(item)}</font>" for item in evidence]),
            info_box("IP PROTECTION POSITION", ip_note, PALE_CYAN),
        ]
    )


def add_applications(story):
    app_profile(
        story, "4", "Bloodchain Demo — public ecosystem and partner narrative", "/", "Investors, prospective partners and public-facing audiences",
        "A single-page, editorial product narrative and launcher that frames the Bloodchain problem, architecture, roadmap and connected applications.",
        [
            "Hero/mission sequence, live national-stat display, product architecture narrative and roadmap treatment.",
            "Expandable roadmap items that open embedded product experiences; mobile application download/link presentation.",
            "Custom ECG-style BLOODCHAIN mark, Botswana/ledger/architecture imagery, dark editorial visual system and Framer Motion reveal/accordion behavior.",
            "Founder/contact call-to-action, a local demo-access modal and local media/lightbox state.",
        ],
        [
            "Read current national statistics; explore system architecture and ledger story; inspect product roadmap; open an embedded product experience or mobile link; start a partner contact action.",
        ],
        "Uses the national-stats API hook for live summary values. The remaining narrative, imagery and interaction structure are local presentation content.",
        [
            "artifacts/demo/src/App.tsx",
            "artifacts/demo/src/pages/home.tsx",
            "artifacts/demo/src/components/*",
            "artifacts/demo/src/index.css",
            "artifacts/demo/public/botswana-map.jpg",
            "artifacts/demo/public/ledger-abstract.jpg",
            "artifacts/demo/public/sovereign-architecture.jpg",
        ],
        "Claim the authored narrative, information architecture, product composition, bespoke graphical mark, custom interaction composition and original media where ownership/permission is documented. Do not claim generic animation libraries, React or stock/third-party assets without a verified license record.",
    )
    story.append(PageBreak())
    app_profile(
        story, "5", "Pulse Web — donor portal", "/pulse/", "Blood donors",
        "A web portal for donor discovery, registration, personal donor records, education and facility discovery.",
        [
            "Routes for landing, registration, dashboard, knowledge, facilities and a not-found state.",
            "Clinical/profile registration form, donor-ID handoff, donor dashboard and donation-history filter.",
            "Browser donor-card download flow, eligibility/compatibility education, facility listing and persistent demo panel.",
            "Warm cream/red donor-facing visual identity distinct from operational applications.",
        ],
        [
            "View supply context; register a donor; carry the returned donor ID into the dashboard; retrieve donor profile/history; filter history; download a donor card; discover facilities and review donation guidance.",
        ],
        "Consumes generated React Query hooks for national statistics, donor registration/retrieval, blood-unit history and facilities. It uses the shared OpenAPI client rather than duplicating server contracts.",
        [
            "artifacts/pulse/src/App.tsx",
            "artifacts/pulse/src/pages/home.tsx",
            "artifacts/pulse/src/pages/register.tsx",
            "artifacts/pulse/src/pages/dashboard.tsx",
            "artifacts/pulse/src/pages/knowledge.tsx",
            "artifacts/pulse/src/pages/facilities.tsx",
            "artifacts/pulse/src/components/*",
            "artifacts/pulse/src/index.css",
        ],
        "The claimable value is the donor workflow, custom page composition, educational framing, donor-card experience and Bloodchain-specific visual expression. The common UI primitives and React Query machinery are third-party/configured components.",
    )
    story.append(PageBreak())
    app_profile(
        story, "6", "Pulse Mobile — native donor companion", "Expo / Android build", "Blood donors using mobile devices",
        "A React Native/Expo donor experience that carries the Pulse flow into a native tabbed application.",
        [
            "Home, Register and My Card/Dashboard tab experiences, plus root layout and not-found handling.",
            "Persistent donor-ID context backed by AsyncStorage; refresh and clear/start-over paths.",
            "Native sharing/ViewShot donor card export and a custom cross-browser PNG download implementation for web use.",
            "Platform-specific tab presentation: iOS NativeTabs liquid-glass path and Android/web blur treatment.",
        ],
        [
            "View national stock context; register with donor details; retain donor identity locally; load personal profile and owned units; share or download the donor card; refresh or reset the local session.",
        ],
        "Uses the same shared API operations as Pulse Web: national stats, donor registration, donor lookup and donor-filtered units. App configuration sets the mobile API origin separately from same-origin web clients.",
        [
            "artifacts/pulse-mobile/app/_layout.tsx",
            "artifacts/pulse-mobile/app/(tabs)/_layout.tsx",
            "artifacts/pulse-mobile/app/(tabs)/index.tsx",
            "artifacts/pulse-mobile/app/(tabs)/register.tsx",
            "artifacts/pulse-mobile/app/(tabs)/dashboard.tsx",
            "artifacts/pulse-mobile/lib/donor-context.tsx",
            "artifacts/pulse-mobile/lib/donor-card-download.ts",
            "artifacts/pulse-mobile/app.json",
        ],
        "The original work includes the mobile flow, persistence pattern, donor card/export experience and platform-specific presentation choices. Expo/React Native, native modules and EAS are third-party/platform infrastructure.",
    )
    story.append(PageBreak())
    app_profile(
        story, "7", "Sanctum — hospital blood-bank workstation", "/sanctum/", "Hospital blood-bank technicians and facility operations staff",
        "A facility-focused terminal for tracking local inventory, receiving/logging units, viewing ledger information and requesting supply routing.",
        [
            "Routes for overview, log unit, transfers, ledger and not-found.",
            "Selected-facility context, inventory and national-status display, facility-unit intake form and local ledger filtering.",
            "Shortage and facility discovery view with a request-routing action.",
            "Dark scanline/command-console visual system, facility badge and blue-green operational language.",
        ],
        [
            "Select a facility; inspect inventory/national context; log a blood unit; identify shortage/facility availability; submit a route request; view the ledger-filtered unit list.",
        ],
        "Consumes facility inventory, national statistics, ledger summary, facility discovery, blood-unit logging, national-shortage and route-request operations from the shared client.",
        [
            "artifacts/sanctum/src/App.tsx",
            "artifacts/sanctum/src/components/layout.tsx",
            "artifacts/sanctum/src/pages/overview.tsx",
            "artifacts/sanctum/src/pages/log-unit.tsx",
            "artifacts/sanctum/src/pages/transfers.tsx",
            "artifacts/sanctum/src/pages/ledger.tsx",
            "artifacts/sanctum/src/index.css",
        ],
        "The claimable contribution is the facility workflow, custom terminal composition, local inventory/transfer experience and Bloodchain-specific interaction design. It is not a claim to generic dashboard widgets or underlying UI libraries.",
    )
    story.append(PageBreak())
    app_profile(
        story, "8", "Vigil — national command dashboard", "/vigil/", "National/ministry coordinators and operational oversight teams",
        "A national-level command view that turns facility, inventory and shortage data into operational awareness and supply-routing support.",
        [
            "Routes for overview, facility visibility, alerts/routing and not-found.",
            "National inventory/shortage/ledger panels with distinct periodic refresh intervals.",
            "Facility search and expandable facility inventory; severity-coded shortage alerts; route calculator.",
            "Ministry-scale dark telemetry style with charts, live clock and operational panels.",
        ],
        [
            "Monitor national stock and deficit signals; inspect a facility; search/expand inventory; review alerts; create a compatible route recommendation for a blood request.",
        ],
        "Consumes national stats, national shortage, ledger summary, facilities, facility inventory and route-request operations. The client refresh cadence is implemented in the Vigil page layer.",
        [
            "artifacts/vigil/src/App.tsx",
            "artifacts/vigil/src/components/layout.tsx",
            "artifacts/vigil/src/pages/Overview.tsx",
            "artifacts/vigil/src/pages/Facilities.tsx",
            "artifacts/vigil/src/pages/Alerts.tsx",
            "artifacts/vigil/src/index.css",
        ],
        "The claimable value lies in the national coordination product composition, operational telemetry framing, facility/shortage workflow and Bloodchain-specific decision presentation. Underlying chart/UI libraries must remain attributed to their respective licenses.",
    )
    story.append(PageBreak())
    app_profile(
        story, "9", "Crucible — blood-processing laboratory workstation", "/crucible/", "Blood-processing laboratory personnel",
        "A dedicated laboratory application that turns collected units into screened and separated components while enforcing a defined processing state model and event trail.",
        [
            "Dashboard, searchable/filterable worklist, unit detail and laboratory ledger routes.",
            "Screening record for ABO/Rh and HIV, Hepatitis B/C, malaria and syphilis markers.",
            "Component separation for red cells, plasma and platelets; processing/quarantine/release/discard actions.",
            "Stage and risk badges, terminal/expired read-only behavior, mutation feedback and event-history display.",
            "Dense dark laboratory-instrument visual language with monospace identifiers and high-contrast status coding.",
        ],
        [
            "Open the worklist; filter by stage/risk; inspect a unit; record screening; move an eligible unit into processing; create components; quarantine/release/discard subject to rules; review chronological audit events and chain hashes.",
        ],
        "Consumes the dedicated laboratory API: dashboard, worklist, unit detail, screening submission, component separation, transition and event stream. Client-side mutations invalidate/reload dashboard, worklist, detail and ledger data.",
        [
            "artifacts/crucible/src/App.tsx",
            "artifacts/crucible/src/components/layout.tsx",
            "artifacts/crucible/src/components/badges.tsx",
            "artifacts/crucible/src/pages/dashboard.tsx",
            "artifacts/crucible/src/pages/worklist.tsx",
            "artifacts/crucible/src/pages/unit-detail.tsx",
            "artifacts/crucible/src/pages/ledger.tsx",
            "artifacts/crucible/src/index.css",
            "screenshots/crucible-dashboard.jpg",
        ],
        "This is one of the strongest standalone product assets: its laboratory workflow, screening/component model, state-guarded transitions, event-chain presentation and distinct workstation expression are concrete original implementation areas. It should be protected with careful contributor assignment and source-history preservation.",
    )
    story.append(PageBreak())
    app_profile(
        story, "10", "Bloodchain Pitch — interactive slide deck", "/bloodchain-pitch/", "Investors, partners and presentation audiences",
        "A ten-slide pitch artifact with authored narrative, fixed-canvas slide components and a browser presentation viewer.",
        [
            "Ten individual slide components, a slide manifest, contact-sheet route and viewer route.",
            "Keyboard, click, touch and parent-frame navigation; a fixed 1920×1080 design contract adapted to responsive viewing.",
            "Slides covering problem framing, system narrative, partnership and build call-to-action.",
            "Slide data/position tooling for editor validation and repair.",
        ],
        [
            "Start the viewer; advance/reverse slides using keyboard, click or touch; open individual slides or all-slides contact sheet; allow a host frame to control navigation through messages.",
        ],
        "It has no operational API dependency. The narrative, slide layout/components, viewer behavior and presentation styling are authored content.",
        [
            "artifacts/bloodchain-pitch/src/App.tsx",
            "artifacts/bloodchain-pitch/src/slideLoader.ts",
            "artifacts/bloodchain-pitch/src/pages/slides/Slide1.tsx through Slide10.tsx",
            "artifacts/bloodchain-pitch/src/data/slides-manifest.json",
            "artifacts/bloodchain-pitch/src/index.css",
            "artifacts/bloodchain-pitch/scripts/*",
        ],
        "Protect the original storyline, visual composition, slide-specific copy and interactive viewer composition. The presentation framework and any third-party image/fonts require their own attribution or rights review.",
    )
    story.append(PageBreak())
    app_profile(
        story, "11", "Mockup Sandbox — internal component preview environment", "Internal artifact / component preview server", "Bloodchain design and implementation workflow",
        "A development-only component-preview surface used to render, compare and iterate on UI components without making it part of the end-user Bloodchain product.",
        [
            "Vite-based preview shell, mockup preview plugin and source application/styles.",
            "Artifact registration as a design/canvas workspace and a dedicated preview-server workflow.",
        ],
        [
            "Host isolated component previews during design iteration; use as supporting evidence of the team's design/development process rather than as a public product module.",
        ],
        "This artifact does not represent a business-user workflow. It supports the production applications but should be categorised separately in commercial and IP discussions.",
        [
            "artifacts/mockup-sandbox/src/App.tsx",
            "artifacts/mockup-sandbox/src/index.css",
            "artifacts/mockup-sandbox/mockupPreviewPlugin.ts",
            "artifacts/mockup-sandbox/.replit-artifact/artifact.toml",
        ],
        "Claim only any original preview orchestration or components actually authored here. Do not overstate it as a client-facing Bloodchain module; it is internal tooling.",
    )
    story.append(PageBreak())


def add_backend_and_data(story):
    story.extend(
        [
            p("12. Shared API, domain contract and database", "h1"),
            p(
                "The shared back end is the operational backbone for every product surface. It is an Express service "
                "backed by Drizzle/PostgreSQL. The HTTP interface is declared once in an OpenAPI specification, "
                "then used to generate typed React Query hooks and Zod validators. This structure is valuable both "
                "technically and as evidence: it makes the relationship between domain rules, contract and clients "
                "traceable.",
                "body",
            ),
            p("API operation inventory", "h2"),
            data_table(
                ["Domain", "Operations under the API root", "Implemented purpose"],
                [
                    ("Health", "GET /health", "Service health/status response."),
                    ("Donors", "GET /donors; POST /donors; GET /donors/{donorId}", "List/filter donors, register a donor, retrieve a donor record."),
                    ("Blood units", "GET /units; POST /units; GET /units/{unitId}", "List/filter units, log a unit and retrieve a unit."),
                    ("Facilities", "GET /facilities; GET /facilities/{facilityId}; GET /facilities/{facilityId}/inventory", "Find facility records and inspect local inventory/critical categories."),
                    ("Routing", "POST /route; GET /route/national-shortage", "Rank available supply candidates and summarise shortage against threshold data."),
                    ("National statistics", "GET /stats/national; GET /stats/ledger", "Aggregate national operational metrics and a presentation-oriented ledger summary."),
                    ("Laboratory", "GET /lab/worklist; GET /lab/units/{unitId}; POST /lab/units/{unitId}/screening; POST /lab/units/{unitId}/components; POST /lab/units/{unitId}/transition; GET /lab/dashboard; GET /lab/events", "Serve and mutate the dedicated processing worklist, laboratory unit detail, component records, state transitions and event history."),
                ],
                [3.0 * cm, 8.0 * cm, 7.8 * cm],
            ),
            p("Core data entities", "h2"),
            data_table(
                ["Entity", "Key responsibility", "Selected content / relationship"],
                [
                    ("donors", "Donor identity and eligibility context", "Donor ID, contact/profile information, blood type, district, totals, status and eligibility dates."),
                    ("blood_units", "Collected unit lifecycle", "Unit ID, donor/facility references, blood type, status, collection/expiry time, temperature and a base chain hash."),
                    ("facilities", "Networked storage/threshold context", "Facility identity, district, type/status, JSON inventory and threshold values."),
                    ("lab_processing", "One laboratory processing record per unit", "Facility, stage, risk, screening JSON, event sequence, chain head and timestamps."),
                    ("lab_components", "Separated component records", "Component identity, parent unit, type, volume, status and expiry."),
                    ("lab_events", "Chronological laboratory event trail", "Event ID, unit/facility, action, actor, reason, sequence, timestamp and hash; indexed for unit/facility/time retrieval."),
                ],
                [3.1 * cm, 5.3 * cm, 10.4 * cm],
            ),
            p("Evidence locations", "h2"),
            bullet_list(
                [
                    "<font name='Courier'>lib/api-spec/openapi.yaml</font> — authored contract and schema source.",
                    "<font name='Courier'>artifacts/api-server/src/routes/index.ts</font> and route modules — endpoint implementation boundary.",
                    "<font name='Courier'>lib/db/src/schema/donors.ts</font>, <font name='Courier'>blood_units.ts</font>, <font name='Courier'>facilities.ts</font>, <font name='Courier'>laboratory.ts</font> — persistent domain model.",
                    "<font name='Courier'>lib/db/src/seed.ts</font>, <font name='Courier'>demo-data.ts</font>, <font name='Courier'>lab-demo-data.ts</font> — deterministic demonstration dataset and operational scenarios.",
                    "<font name='Courier'>lib/api-spec/orval.config.ts</font> — code generation configuration and source/output relationship.",
                ]
            ),
            PageBreak(),
            p("13. Laboratory workflow and chain-of-custody mechanics", "h1"),
            p(
                "Crucible's server-side laboratory rules are a focused source of original implementation. The workflow "
                "does not simply render statuses: it uses transactional mutation logic and row locking to protect "
                "unit processing. The following description records implemented behavior, not a claim of clinical "
                "certification.",
                "body",
            ),
            data_table(
                ["Lifecycle stage", "Implemented transition / guard", "Recorded event behavior"],
                [
                    ("awaiting_tests", "Initial test queue; may move to processing when screening outcome is suitable.", "Event appended with actor, reason/time context and next sequence."),
                    ("processing", "Used for clear screened unit processing; component separation requires completed clear screening.", "Screening and processing actions are retained in the unit event stream."),
                    ("quarantine", "Used for reactive/review results and separated components; release requires cleared screening, components and unexpired products.", "Quarantine requires a reason in the state-transition input."),
                    ("released", "Terminal release state after prerequisites are met; later transitions/re-screening are rejected.", "Release event completes the demonstrable processing trail."),
                    ("discarded", "Terminal state; requires a reason and synchronises relevant status treatment.", "Discard event records reason/actor context."),
                ],
                [3.1 * cm, 9.1 * cm, 6.6 * cm],
            ),
            p("Screening, separation and risk derivation", "h2"),
            bullet_list(
                [
                    "Screening input includes ABO group, Rh factor, HIV, Hepatitis B, Hepatitis C, malaria and syphilis markers.",
                    "Risk is derived from results: positive markers become reactive; pending results become pending; ABO/Rh mismatch becomes review; otherwise the outcome can be clear.",
                    "Component separation validates an eligible clear-processing unit, prevents duplicate component types, limits declared volume and creates red-cell, plasma and platelet component records with distinct shelf-life assumptions.",
                    "Release is blocked unless the implemented prerequisites are satisfied; expired base units or components and terminal/transfused states are rejected.",
                ]
            ),
            p("Event-chain implementation", "h2"),
            p(
                "Each laboratory event is appended in a transaction after the processing record's prior sequence and "
                "chain head are read. The next hash is SHA-256-derived from the prior hash plus the unit, facility, "
                "sequence, action, actor, reason and timestamp context; the processing record is updated with the "
                "new sequence/head. Startup logic also repairs legacy entries missing explicit sequence/head data. "
                "This gives the product a concrete chain-of-custody design to preserve and mature, while not by "
                "itself constituting an externally validated blockchain or legal non-repudiation system.",
                "body",
            ),
            info_box(
                "IMPORTANT CLAIM BOUNDARY",
                "The current national <font name='Courier'>/stats/ledger</font> summary is a presentation-oriented "
                "ledger statistic and is distinct from the laboratory <font name='Courier'>lab_events</font> chain. "
                "When pitching, describe the laboratory event chain precisely; do not imply every national metric "
                "is backed by the same immutable event mechanism.",
            ),
            p("Key evidence locations", "h2"),
            bullet_list(
                [
                    "<font name='Courier'>artifacts/api-server/src/lib/lab-mutations.ts</font> — transactional guards, state transitions, component logic and event append.",
                    "<font name='Courier'>artifacts/api-server/src/lib/lab-service.ts</font> — worklist/detail/metric/event read service and risk derivation.",
                    "<font name='Courier'>lib/db/src/lab-ledger.ts</font> — compatibility repair for legacy event sequence/chain head.",
                    "<font name='Courier'>lib/db/src/schema/laboratory.ts</font> — processing/component/event tables and indexes.",
                ]
            ),
            PageBreak(),
        ]
    )


def add_shared_and_boundary(story):
    story.extend(
        [
            p("14. Shared libraries, build system and operating model", "h1"),
            data_table(
                ["Package / location", "Role", "Ownership and evidence treatment"],
                [
                    ("lib/api-spec", "OpenAPI source document and Orval generation configuration.", "Treat the authored OpenAPI document and its Bloodchain domain schema as key source evidence."),
                    ("lib/api-client-react", "Generated React Query operations plus custom fetch wrapper.", "Generated portions are derivative from the spec; custom wrapper is authored support code."),
                    ("lib/api-zod", "Generated server request/response validators and type exports.", "Generated portions are derivative from the spec; preserve for build reproducibility, but anchor claims in the OpenAPI source."),
                    ("lib/db", "Drizzle connection, schema, sequence helpers, seed/demo datasets, laboratory data and ledger repair.", "High-value authored domain implementation."),
                    ("artifacts/api-server", "Express application, route mounting, logging, lab service/mutations and startup seed/init sequence.", "High-value authored service implementation."),
                    ("pnpm root manifests / artifact TOMLs / scripts", "Workspace, build, artifact routing, production metadata and post-merge helper.", "Operational evidence; platform/tool configuration itself has narrower protectable scope."),
                ],
                [4.1 * cm, 7.0 * cm, 7.7 * cm],
            ),
            p("Runtime and deployment characteristics", "h2"),
            bullet_list(
                [
                    "Private pnpm workspace with Node.js/TypeScript; root build performs library typechecking and package builds.",
                    "Web artifacts use Vite with injected PORT and BASE_PATH, static output and SPA-style route handling.",
                    "The API is Express/Drizzle/PostgreSQL and requires a database connection; startup runs idempotent facility/demo/laboratory initialisation.",
                    "The mobile application is Expo/React Native; EAS profiles support an internal Android APK and a production Android App Bundle.",
                    "Replit artifact metadata registers the delivery surfaces. This is deployment infrastructure, not a substitute for a formal source-control, CI/CD and release-evidence policy.",
                ]
            ),
            p("Third-party and generated boundary", "h2"),
            data_table(
                ["Category", "Examples in the repository", "Recommended treatment in an IP pack"],
                [
                    ("Open-source application stack", "React, React DOM, Vite, TypeScript, Wouter, TanStack React Query, Tailwind, Framer Motion, Lucide, Radix/shadcn-style primitives.", "Retain license notices and package lockfile. Do not claim ownership; claim the Bloodchain-specific orchestration, workflows and presentation."),
                    ("API/data stack", "Express, PostgreSQL/node-postgres, Drizzle ORM, Zod, drizzle-zod, Orval, esbuild, Pino.", "Treat as dependencies. Preserve version/lock evidence and distinguish authored route/schema/mutation code."),
                    ("Mobile/platform stack", "Expo, Expo Router, React Native, EAS and related native modules.", "Treat as third-party platform tooling. Claim native workflow composition and authored application code, not the platform."),
                    ("Replit platform tooling", "Artifact metadata, Replit Vite plugins, connectors SDK, managed workflow/deployment facilities.", "Treat as platform infrastructure; preserve configurations as operational evidence, not as exclusive technology ownership."),
                    ("Generated output", "lib/api-client-react/src/generated, lib/api-zod/src/generated, compiled dist output.", "Preserve for reproducibility but cite OpenAPI and authored source code as primary evidence."),
                ],
                [3.7 * cm, 7.0 * cm, 8.1 * cm],
            ),
            p("15. IP claim map and evidence to preserve", "h1"),
            data_table(
                ["Potential proprietary asset", "Concrete implementation evidence", "Protection / preservation action"],
                [
                    ("Bloodchain product architecture", "Multi-role donor, facility, ministry and lab surfaces joined through a shared API/domain model.", "Keep versioned architecture diagrams, Git history, pitch chronology and dated releases. Describe as a product/system design, not a claim to generic health dashboards."),
                    ("Donor engagement workflows", "Pulse Web/Mobile registration, donor identity handoff, history, card export, facilities and education.", "Preserve source, UX flows, visual comps and records of authorship. Review wording/imagery licenses."),
                    ("Facility allocation and shortage workflows", "Facility inventory/threshold data, national shortage computation, compatible candidate routing and user interfaces.", "Preserve route/routing source, API contract and scenario test evidence. Avoid claiming unverified clinical optimisation performance."),
                    ("Laboratory lifecycle rules", "Screening results, risk derivation, component separation, quarantine/release/discard constraints and terminal protections.", "Preserve lab mutation code, schema, UI and test records. Seek clinical/legal review before production claims."),
                    ("Chain-of-custody event mechanism", "Per-unit sequence, hash head, hash payload, transactional append and ordered ledger display.", "Preserve source history and example event exports; explain its implemented scope accurately."),
                    ("Brand and creative expression", "Bloodchain/Crucible/Sanctum/Vigil/Pulse names, visual language, slides, logo mark, original copy and custom page composition.", "Run trademark/name clearance; keep original source/design files; record ownership or licence of every image/font."),
                    ("OpenAPI/domain vocabulary", "Explicit endpoints and schema names for donors, units, facilities, routing and laboratory processing.", "Archive versioned spec snapshots and generated output mapping; use contributor agreements for future API changes."),
                ],
                [4.0 * cm, 7.2 * cm, 7.6 * cm],
            ),
            PageBreak(),
        ]
    )


def add_risks_and_appendix(story):
    story.extend(
        [
            p("16. Security, data-sharing and production-readiness boundary", "h1"),
            p(
                "This report is about technical authorship and evidence, but safe external pitching requires clear "
                "statements about what is demonstrable today and what still needs institutional governance before "
                "real data or clinical operations are introduced.",
                "body",
            ),
            data_table(
                ["Current observed state", "Implication for a partner discussion", "Recommended response before/alongside a pilot"],
                [
                    ("Demo/seeded operational records are present.", "Do not represent the displayed records as live NBTS clinical or national data.", "Use labelled synthetic/demo datasets; document source/permission for all imagery and scenario data."),
                    ("Crucible has no production authentication or role enforcement in the current implementation; operator name is currently supplied by callers.", "Do not present release/discard actions as ready for operational deployment.", "Add authenticated identities, facility-scoped RBAC, server-derived actor attribution and audit review before real use."),
                    ("Laboratory hardware, EHR/LIS and courier integrations are out of current scope.", "Do not imply automated instrument ingestion or end-to-end clinical interoperability.", "Define interface, verification, exception and accountability requirements with NBTS."),
                    ("Database schema is deployed via an automated push-style build command.", "Change governance and production migrations require more formal control before institutional rollout.", "Adopt migration review, backup/rollback and change-approval practices."),
                    ("The codebase contains third-party packages and generated material.", "Ownership claims need an open-source/commercial-asset boundary.", "Create an SBOM/licence register; retain lockfile and notices; complete image/font provenance review."),
                ],
                [5.1 * cm, 6.2 * cm, 7.5 * cm],
            ),
            p("17. Evidence preservation checklist before pitching", "h1"),
            bullet_list(
                [
                    "<b>Freeze a dated source snapshot.</b> Create a signed/tagged Git commit, archive it in a controlled repository and retain this report alongside the exact commit hash.",
                    "<b>Record authorship and assignment.</b> Keep contributor identity, employment/contractor IP assignment, work dates and review records for code, design, copy, images and decks.",
                    "<b>Hash the evidence package.</b> Generate checksums for source export, PDF report, pitch deck export and key design assets; store the signed/checksummed register separately.",
                    "<b>Preserve build evidence.</b> Archive build logs, production/dev artifact configuration, package lockfile and released application captures.",
                    "<b>Clear names and creative assets.</b> Obtain trademark/name advice for Bloodchain and product names; record the provenance or licence for logo elements, photos, fonts and illustrations.",
                    "<b>Control external disclosure.</b> Use a redacted capability deck for early conversations. Move detailed API/schema/lab-chain explanations into governed technical sessions under suitable confidentiality terms.",
                    "<b>Prepare data-governance material.</b> Before receiving any identifiable health data, agree a written data-sharing framework covering purpose, legal basis, minimum dataset, security, retention, breach handling, ownership, access and deletion/return.",
                ]
            ),
            info_box(
                "LEGAL NOTE",
                "This inventory is evidence-oriented, not a legal registration. Engage qualified Botswana IP, "
                "commercial and data-protection counsel to advise on trademarks, copyright assignment, database "
                "rights, confidentiality, procurement, health-data governance and any pilot agreement before relying "
                "on this report in a legal dispute or regulated deployment.",
            ),
            p("Appendix A — source evidence index", "h1"),
            p(
                "The following index lists the principal authored locations used in this report. It is intentionally "
                "a source map rather than a full file dump; the complete repository and Git history remain the "
                "authoritative granular evidence.",
                "body",
            ),
            data_table(
                ["Area", "Principal source locations", "What the source evidences"],
                [
                    ("Public narrative", "artifacts/demo/src; artifacts/demo/public; artifacts/demo/.replit-artifact/artifact.toml", "Public product narrative, original page composition/assets, deployment surface."),
                    ("Donor web", "artifacts/pulse/src/App.tsx; src/pages; src/components; src/index.css", "Registration/dashboard/knowledge/facility/card experience."),
                    ("Donor mobile", "artifacts/pulse-mobile/app; lib/donor-context.tsx; lib/donor-card-download.ts; app.json", "Native flows, local donor persistence and card export."),
                    ("Hospital operations", "artifacts/sanctum/src/App.tsx; components/layout.tsx; pages; index.css", "Facility inventory, unit intake, transfer and facility ledger surface."),
                    ("National operations", "artifacts/vigil/src/App.tsx; components/layout.tsx; pages; index.css", "National dashboard, shortage, facilities and routing surface."),
                    ("Laboratory operations", "artifacts/crucible/src; artifacts/api-server/src/lib/lab-*; lib/db/src/schema/laboratory.ts", "Screening, processing, component, state and chain-of-custody experience."),
                    ("Shared contract", "lib/api-spec/openapi.yaml; lib/api-spec/orval.config.ts", "Domain interface, request/response types and generation boundary."),
                    ("Shared data", "lib/db/src/schema; seed.ts; demo-data.ts; lab-demo-data.ts; lab-ledger.ts; sequences.ts", "Persistent data model, demo datasets, ID/event support."),
                    ("Server routes", "artifacts/api-server/src/app.ts; index.ts; routes/*; lib/*", "Express application, operational API routes, routing and laboratory business logic."),
                    ("Pitch content", "artifacts/bloodchain-pitch/src/App.tsx; pages/slides/*; slideLoader.ts; data/slides-manifest.json; index.css", "Authored pitch narrative, interaction and slide composition."),
                    ("Build/platform", "package.json; pnpm-workspace.yaml; artifact TOMLs; scripts/post-merge.sh; replit.md", "Workspace/build/deployment context and dependency policy."),
                ],
                [3.0 * cm, 8.5 * cm, 7.3 * cm],
            ),
            p("Appendix B — visual evidence", "h1"),
        ]
    )

    screenshot = ROOT / "screenshots" / "crucible-dashboard.jpg"
    if screenshot.exists():
        img = Image(str(screenshot))
        img._restrictSize(17.2 * cm, 10.5 * cm)
        story.extend(
            [
                p(
                    "Current Crucible operational dashboard capture from the running artifact. This is visual evidence "
                    "of the authored laboratory dashboard and its populated workflow/event presentation; it is not "
                    "evidence of a live clinical deployment.",
                    "small",
                ),
                Spacer(1, 0.15 * cm),
                img,
            ]
        )
    else:
        story.append(p("No local visual capture was available when this report was generated.", "small"))
    story.extend(
        [
            Spacer(1, 0.4 * cm),
            p("End of confidential internal report", "callout"),
        ]
    )


def build_report():
    DELIVERABLES.mkdir(exist_ok=True)
    doc = SimpleDocTemplate(
        str(REPORT_PATH),
        pagesize=A4,
        leftMargin=1.55 * cm,
        rightMargin=1.55 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.25 * cm,
        title="Bloodchain Confidential Codebase & Intellectual Property Protection Report",
        author="Bloodchain",
        subject="Internal technical inventory and IP evidence map",
    )
    story = []
    cover(story)
    add_report_intro(story)
    add_architecture(story)
    add_applications(story)
    add_backend_and_data(story)
    add_shared_and_boundary(story)
    add_risks_and_appendix(story)
    doc.build(story, onFirstPage=report_header, onLaterPages=report_header)


def build_letter():
    DELIVERABLES.mkdir(exist_ok=True)
    doc = SimpleDocTemplate(
        str(LETTER_PATH),
        pagesize=A4,
        leftMargin=2.05 * cm,
        rightMargin=2.05 * cm,
        topMargin=1.55 * cm,
        bottomMargin=1.5 * cm,
        title="Draft Introduction Letter — Bloodchain to National Blood Transfusion Service",
        author="Bloodchain",
        subject="Request for introductory meeting and data-sharing discussion",
    )
    story = [
        Spacer(1, 0.35 * cm),
        p("BLOODCHAIN", "h3"),
        p("Draft introduction letter", "letter_title"),
        p(
            "<b>To:</b> The Director<br/>"
            "National Blood Transfusion Service<br/>"
            "Ministry of Health and Wellness<br/>"
            "Republic of Botswana",
            "letter",
        ),
        p(f"<b>Date:</b> {PREPARED_DATE}", "letter"),
        p("<b>Subject:</b> Request for an introductory meeting on Bloodchain and a staged data-sharing framework", "letter"),
        p("Dear Director,", "letter"),
        p(
            "I am writing on behalf of Bloodchain, a Botswana-focused digital health initiative exploring how a "
            "connected blood-supply platform could support visibility, coordination and accountable workflows across "
            "donors, hospital blood banks, national operations and blood-processing laboratories.",
            "letter",
        ),
        p(
            "We have developed an early product demonstration comprising a donor portal and mobile companion, a "
            "hospital blood-bank workstation, a national operations dashboard and a laboratory processing workspace. "
            "The platform is designed to make the path from donation through processing, local inventory, shortage "
            "identification and facility coordination easier to understand and manage. The current build is a "
            "demonstration environment; it is not presented as a replacement for NBTS systems, a clinically certified "
            "production system, or a request for unrestricted access to NBTS data.",
            "letter",
        ),
        p(
            "We would value the opportunity to meet with the National Blood Transfusion Service to understand your "
            "strategic priorities, current information flows, governance requirements and the practical problems "
            "that a locally grounded digital solution should address. Our first objective would be to listen, "
            "demonstrate the concept at an appropriate level, and identify whether a small, jointly governed discovery "
            "or pilot pathway could be useful.",
            "letter",
        ),
        p(
            "If there is interest, we would propose that any future data-sharing discussion be staged and documented "
            "through the appropriate NBTS, Ministry, legal, privacy, information-security and clinical-governance "
            "channels. A framework could cover:",
            "letter",
        ),
        bullet_list(
            [
                "a clearly defined purpose, scope, legal basis and governance/approval pathway;",
                "the minimum data required for an initial feasibility exercise, with de-identification or synthetic data used wherever possible;",
                "role-based access, auditability, encryption, hosting, retention, incident handling and return/deletion requirements;",
                "data quality, validation, clinical accountability and the limits of any decision-support functionality;",
                "ownership of pre-existing intellectual property, treatment of jointly developed outputs, confidentiality and publication/communications approval;",
                "a time-bound pilot plan with defined success measures, review checkpoints and an exit process.",
            ]
        ),
        p(
            "We would be grateful for a 45–60 minute introductory meeting at a time convenient to you. We can provide "
            "a concise non-technical overview in advance and, subject to your preference, a separate technical "
            "briefing for the relevant operational and information-governance teams. We are committed to approaching "
            "any collaboration respectfully, transparently and in alignment with Botswana’s health-sector and "
            "data-protection requirements.",
            "letter",
        ),
        p(
            "Thank you for considering this request. We would welcome the opportunity to learn from the National "
            "Blood Transfusion Service and explore whether Bloodchain could be shaped into a responsible, locally "
            "useful complement to existing national efforts.",
            "letter",
        ),
        Spacer(1, 0.2 * cm),
        p("Yours faithfully,", "letter"),
        Spacer(1, 0.6 * cm),
        p(
            "<b>[Your full name]</b><br/>"
            "[Your title / Founder, Bloodchain]<br/>"
            "[Email address]  |  [Telephone number]<br/>"
            "[Optional website or presentation link]",
            "letter",
        ),
        Spacer(1, 0.4 * cm),
        info_box(
            "DRAFTING NOTE",
            "Replace the bracketed sender details and confirm the addressee/title with NBTS before sending. Have "
            "the final version reviewed by appropriate legal, privacy and clinical-governance advisers before it "
            "creates any data-sharing commitment.",
            PALE_CYAN,
        ),
    ]
    doc.build(story, onFirstPage=letter_header, onLaterPages=letter_header)


def build_letter_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(11, 18, 32)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.12

    if "Letter Subtitle" not in [style.name for style in doc.styles]:
        subtitle = doc.styles.add_style("Letter Subtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle.base_style = doc.styles["Normal"]
        subtitle.font.name = "Arial"
        subtitle.font.size = Pt(18)
        subtitle.font.bold = True
        subtitle.font.color.rgb = RGBColor(11, 18, 32)
        subtitle.paragraph_format.space_after = Pt(14)

    header = section.header.paragraphs[0]
    header.text = "BLOODCHAIN"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.bold = True
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(122, 16, 37)

    footer = section.footer.paragraphs[0]
    footer.text = "Draft for review • Replace bracketed fields before sending"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    title = doc.add_paragraph(style="Letter Subtitle")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Draft introduction letter")

    for line in [
        "To: The Director",
        "National Blood Transfusion Service",
        "Ministry of Health and Wellness",
        "Republic of Botswana",
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph(f"Date: {PREPARED_DATE}")
    subject = doc.add_paragraph()
    subject.add_run("Subject: ").bold = True
    subject.add_run(
        "Request for an introductory meeting on Bloodchain and a staged data-sharing framework"
    )
    doc.add_paragraph("Dear Director,")

    paragraphs = [
        (
            "I am writing on behalf of Bloodchain, a Botswana-focused digital health initiative "
            "exploring how a connected blood-supply platform could support visibility, coordination "
            "and accountable workflows across donors, hospital blood banks, national operations and "
            "blood-processing laboratories."
        ),
        (
            "We have developed an early product demonstration comprising a donor portal and mobile "
            "companion, a hospital blood-bank workstation, a national operations dashboard and a "
            "laboratory processing workspace. The platform is designed to make the path from donation "
            "through processing, local inventory, shortage identification and facility coordination "
            "easier to understand and manage. The current build is a demonstration environment; it is "
            "not presented as a replacement for NBTS systems, a clinically certified production system, "
            "or a request for unrestricted access to NBTS data."
        ),
        (
            "We would value the opportunity to meet with the National Blood Transfusion Service to "
            "understand your strategic priorities, current information flows, governance requirements "
            "and the practical problems that a locally grounded digital solution should address. Our "
            "first objective would be to listen, demonstrate the concept at an appropriate level, and "
            "identify whether a small, jointly governed discovery or pilot pathway could be useful."
        ),
        (
            "If there is interest, we would propose that any future data-sharing discussion be staged "
            "and documented through the appropriate NBTS, Ministry, legal, privacy, information-security "
            "and clinical-governance channels. A framework could cover:"
        ),
    ]
    for text in paragraphs:
        doc.add_paragraph(text)

    data_sharing_items = [
        "a clearly defined purpose, scope, legal basis and governance/approval pathway;",
        "the minimum data required for an initial feasibility exercise, with de-identification or synthetic data used wherever possible;",
        "role-based access, auditability, encryption, hosting, retention, incident handling and return/deletion requirements;",
        "data quality, validation, clinical accountability and the limits of any decision-support functionality;",
        "ownership of pre-existing intellectual property, treatment of jointly developed outputs, confidentiality and publication/communications approval;",
        "a time-bound pilot plan with defined success measures, review checkpoints and an exit process.",
    ]
    for item in data_sharing_items:
        doc.add_paragraph(item, style="List Bullet")

    for text in [
        (
            "We would be grateful for a 45–60 minute introductory meeting at a time convenient to you. "
            "We can provide a concise non-technical overview in advance and, subject to your preference, "
            "a separate technical briefing for the relevant operational and information-governance teams. "
            "We are committed to approaching any collaboration respectfully, transparently and in alignment "
            "with Botswana’s health-sector and data-protection requirements."
        ),
        (
            "Thank you for considering this request. We would welcome the opportunity to learn from the "
            "National Blood Transfusion Service and explore whether Bloodchain could be shaped into a "
            "responsible, locally useful complement to existing national efforts."
        ),
        "Yours faithfully,",
    ]:
        doc.add_paragraph(text)

    doc.add_paragraph("\n[Your full name]\n[Your title / Founder, Bloodchain]\n[Email address]  |  [Telephone number]\n[Optional website or presentation link]")

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    run = note.add_run(
        "DRAFTING NOTE: Replace the bracketed sender details and confirm the addressee/title with NBTS "
        "before sending. Have the final version reviewed by appropriate legal, privacy and clinical-governance "
        "advisers before it creates any data-sharing commitment."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(122, 16, 37)
    doc.save(LETTER_DOCX_PATH)


if __name__ == "__main__":
    build_report()
    build_letter()
    build_letter_docx()
    print(f"Created: {REPORT_PATH.relative_to(ROOT)}")
    print(f"Created: {LETTER_PATH.relative_to(ROOT)}")
    print(f"Created: {LETTER_DOCX_PATH.relative_to(ROOT)}")